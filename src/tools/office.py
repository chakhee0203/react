import os
import re
import io
import base64
import uuid
import pandas as pd
from langchain_core.tools import tool
import subprocess
import shutil
import zipfile
import xml.etree.ElementTree as ET
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns

# Configure fonts for Chinese support
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun', 'Arial', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

def _allowed(filename: str) -> bool:
    allowed = os.environ.get("CURRENT_SESSION_UPLOADS", "")
    names = [x.strip() for x in re.split(r"[;,]", allowed) if x.strip()]
    return filename in names

ARTIFACT_CACHE = {}

def _put_artifact(data: str) -> str:
    k = uuid.uuid4().hex
    ARTIFACT_CACHE[k] = data
    return k

def get_artifact(key: str) -> str:
    return ARTIFACT_CACHE.get(key, "")

@tool
def excel_to_csv_from_upload(filename: str, return_base64: bool = False) -> str:
    """Convert an uploaded Excel file to CSV and return content or Base64.
    The file must exist in 'uploads/' directory.
    """
    if not _allowed(filename):
        return "Error: File not allowed (not in current session uploads)"
    path = os.path.join("uploads", filename)
    if not os.path.exists(path):
        return "Error: File not found in uploads/"
    try:
        df = pd.read_excel(path)
        csv_str = df.to_csv(index=False)
        if return_base64:
            b64 = base64.b64encode(csv_str.encode('utf-8')).decode('utf-8')
            aid = _put_artifact(b64)
            return f"Converted successfully. [FILE_ID:{aid}:csv:{os.path.splitext(filename)[0]}.csv]"
        return csv_str
    except Exception as e:
        return f"Error converting Excel to CSV: {str(e)}"
    finally:
        try:
            os.remove(path)
        except Exception:
            pass

def _load_table_from_upload(filename: str):
    path = os.path.join("uploads", filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(path)
    elif ext == ".csv":
        try:
            df = pd.read_csv(path)
        except UnicodeDecodeError:
            df = pd.read_csv(path, encoding="gbk")
    else:
        raise RuntimeError("Unsupported file type. Please upload Excel or CSV.")
    return df, path

@tool
def table_basic_profile_from_upload(filename: str) -> str:
    """Generate a basic profile (dtypes, missing, unique, numeric stats) for an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        rows, cols = df.shape
        dtypes = df.dtypes.astype(str).to_dict()
        miss = df.isna().sum().to_dict()
        uniq = {c: int(df[c].nunique()) for c in df.columns}
        nums = df.select_dtypes(include=["number"])
        desc = nums.describe().to_csv() if not nums.empty else "No numeric columns."
        
        out = []
        out.append(f"Rows: {rows}, Cols: {cols}")
        out.append("\nData Types:")
        for k, v in dtypes.items():
            out.append(f"- {k}: {v}")
        out.append("\nMissing Values:")
        for k, v in miss.items():
            if v > 0:
                out.append(f"- {k}: {v}")
        if all(v == 0 for v in miss.values()):
            out.append("No missing values.")
            
        out.append("\nUnique Values Count:")
        for k, v in uniq.items():
            out.append(f"- {k}: {v}")
            
        # Save numeric summary as artifact
        aid = _put_artifact(base64.b64encode(desc.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + "_numeric_summary.csv"
        
        return "\n".join(out) + f"\n\n[FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error profiling table: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_value_counts_from_upload(filename: str, column: str) -> str:
    """Get value counts for a specific column in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        if column not in df.columns:
            return f"Error: Column '{column}' not found. Available columns: {list(df.columns)}"
            
        vc = df[column].value_counts().to_csv()
        aid = _put_artifact(base64.b64encode(vc.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + f"_{column}_counts.csv"
        
        top_5 = df[column].value_counts().head(5).to_string()
        return f"Top 5 values for '{column}':\n{top_5}\n\nFull counts available: [FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error getting value counts: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_correlation_from_upload(filename: str) -> str:
    """Calculate correlation matrix for numeric columns in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        nums = df.select_dtypes(include=["number"])
        if nums.empty:
            return "Error: No numeric columns found for correlation analysis."
            
        corr = nums.corr()
        csv_out = corr.to_csv()
        aid = _put_artifact(base64.b64encode(csv_out.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + "_correlation.csv"
        
        return f"Correlation matrix calculated. [FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error calculating correlation: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_filter_query_from_upload(filename: str, query: str) -> str:
    """Filter rows in uploaded Excel/CSV using a pandas query string (e.g. 'age > 30')."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        try:
            filtered = df.query(query)
        except Exception as qe:
            return f"Error executing query '{query}': {str(qe)}"
            
        csv_out = filtered.to_csv(index=False)
        aid = _put_artifact(base64.b64encode(csv_out.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + "_filtered.csv"
        
        return f"Filtered {len(filtered)} rows (from {len(df)}). [FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error filtering table: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_outliers_from_upload(filename: str, column: str) -> str:
    """Detect outliers in a numeric column using IQR method (1.5 * IQR)."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        if column not in df.columns:
            return f"Error: Column '{column}' not found."
        if not pd.api.types.is_numeric_dtype(df[column]):
            return f"Error: Column '{column}' is not numeric."
            
        Q1 = df[column].quantile(0.25)
        Q3 = df[column].quantile(0.75)
        IQR = Q3 - Q1
        lower = Q1 - 1.5 * IQR
        upper = Q3 + 1.5 * IQR
        
        outliers = df[(df[column] < lower) | (df[column] > upper)]
        
        csv_out = outliers.to_csv(index=False)
        aid = _put_artifact(base64.b64encode(csv_out.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + f"_{column}_outliers.csv"
        
        return f"Found {len(outliers)} outliers in '{column}' (bounds: {lower:.2f}, {upper:.2f}). [FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error detecting outliers: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_pivot_from_upload(filename: str, index: str, columns: str, values: str, aggfunc: str = "mean") -> str:
    """Create a pivot table from uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        pivot = df.pivot_table(index=index, columns=columns, values=values, aggfunc=aggfunc)
        
        csv_out = pivot.to_csv()
        aid = _put_artifact(base64.b64encode(csv_out.encode("utf-8")).decode("utf-8"))
        name = os.path.splitext(filename)[0] + "_pivot.csv"
        
        return f"Pivot table created. [FILE_ID:{aid}:csv:{name}]"
    except Exception as e:
        return f"Error creating pivot table: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

def _save_plot_to_artifact(filename_prefix: str, chart_type: str = "chart") -> str:
    """Save current matplotlib figure to base64 artifact and return formatted string (display + download)."""
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight')
    plt.close()
    buf.seek(0)
    b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
    aid = _put_artifact(b64)
    
    # Generate a safe filename for download
    safe_name = os.path.splitext(filename_prefix)[0]
    out_name = f"{safe_name}_{chart_type}.png"
    
    # Return raw base64 as requested by user
    return f"[IMAGE:png:{b64}]\n\nDownload Chart: [FILE:png:{b64}:{out_name}]"

@tool
def table_chart_histogram_from_upload(filename: str, column: str, bins: int = 10) -> str:
    """Generate a histogram for a numeric column in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        if column not in df.columns:
            return f"Error: Column '{column}' not found."
        if not pd.api.types.is_numeric_dtype(df[column]):
            return f"Error: Column '{column}' is not numeric."
            
        plt.figure(figsize=(10, 6))
        sns.histplot(df[column], bins=bins, kde=True)
        plt.title(f"Histogram of {column}")
        plt.xlabel(column)
        plt.ylabel("Frequency")
        
        return f"Histogram created. {_save_plot_to_artifact(filename, 'histogram')}"
    except Exception as e:
        return f"Error generating histogram: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_chart_scatter_from_upload(filename: str, x_column: str, y_column: str) -> str:
    """Generate a scatter plot for two numeric columns in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        for col in [x_column, y_column]:
            if col not in df.columns:
                return f"Error: Column '{col}' not found."
            if not pd.api.types.is_numeric_dtype(df[col]):
                return f"Error: Column '{col}' is not numeric."
                
        plt.figure(figsize=(10, 6))
        sns.scatterplot(data=df, x=x_column, y=y_column)
        plt.title(f"Scatter Plot: {x_column} vs {y_column}")
        
        return f"Scatter plot created. {_save_plot_to_artifact(filename, 'scatter')}"
    except Exception as e:
        return f"Error generating scatter plot: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_chart_line_from_upload(filename: str, x_column: str, y_column: str) -> str:
    """Generate a line chart (e.g. time series) for two columns in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        if x_column not in df.columns or y_column not in df.columns:
            return f"Error: Columns not found."
            
        # Try to parse x_column as datetime if it looks like time
        try:
            df[x_column] = pd.to_datetime(df[x_column])
        except Exception:
            pass
            
        plt.figure(figsize=(12, 6))
        sns.lineplot(data=df, x=x_column, y=y_column)
        plt.title(f"Line Chart: {y_column} over {x_column}")
        plt.xticks(rotation=45)
        
        return f"Line chart created. {_save_plot_to_artifact(filename, 'line_chart')}"
    except Exception as e:
        return f"Error generating line chart: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def table_chart_bar_from_upload(filename: str, x_column: str, y_column: str, aggregation: str = "sum") -> str:
    """Generate a bar chart for categorical x and numeric y in an uploaded Excel/CSV."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        df, path = _load_table_from_upload(filename)
        
        if x_column not in df.columns or y_column not in df.columns:
            return f"Error: Columns not found."
        if not pd.api.types.is_numeric_dtype(df[y_column]):
            return f"Error: Column '{y_column}' must be numeric."
            
        # Aggregate data
        if aggregation == "sum":
            data = df.groupby(x_column)[y_column].sum().reset_index()
        elif aggregation == "mean":
            data = df.groupby(x_column)[y_column].mean().reset_index()
        elif aggregation == "count":
            data = df.groupby(x_column)[y_column].count().reset_index()
        else:
            return "Error: Aggregation must be 'sum', 'mean', or 'count'."
            
        # Sort by value descending for better visualization
        data = data.sort_values(y_column, ascending=False).head(20) # Limit to top 20
        
        plt.figure(figsize=(12, 6))
        sns.barplot(data=data, x=x_column, y=y_column)
        plt.title(f"Bar Chart: {y_column} by {x_column} ({aggregation})")
        plt.xticks(rotation=45)
        
        return f"Bar chart created. {_save_plot_to_artifact(filename, 'bar_chart')}"
    except Exception as e:
        return f"Error generating bar chart: {str(e)}"
    finally:
        try:
            os.remove(os.path.join("uploads", filename))
        except Exception:
            pass

@tool
def csv_to_excel_from_upload(filename: str) -> str:
    """Convert an uploaded CSV file to Excel and return Base64 for download."""
    if not _allowed(filename):
        return "Error: File not allowed (not in current session uploads)"
    path = os.path.join("uploads", filename)
    if not os.path.exists(path):
        return "Error: File not found in uploads/"
    try:
        df = pd.read_csv(path)
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        aid = _put_artifact(b64)
        return f"Converted successfully. [FILE_ID:{aid}:xlsx:{os.path.splitext(filename)[0]}.xlsx]"
    except Exception as e:
        return f"Error converting CSV to Excel: {str(e)}"
    finally:
        try:
            os.remove(path)
        except Exception:
            pass

@tool
def markdown_to_html(md_text: str) -> str:
    """Convert Markdown text to HTML."""
    try:
        import markdown as md
        html = md.markdown(md_text)
        return html
    except Exception as e:
        return f"Error converting markdown: {str(e)}"

@tool
def word_to_pdf_from_upload(filename: str) -> str:
    """Convert an uploaded Word (DOCX) file to PDF; returns a downloadable artifact."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        path = os.path.join("uploads", filename)
        if not os.path.exists(path):
            return "Error: File not found in uploads/"
        pdf_bytes = None
        # 1) docx2pdf
        try:
            from docx2pdf import convert as docx2pdf_convert
            out_pdf = os.path.splitext(path)[0] + ".pdf"
            docx2pdf_convert(path, out_pdf)
            with open(out_pdf, "rb") as f:
                pdf_bytes = f.read()
            try:
                os.remove(out_pdf)
            except Exception:
                pass
        except Exception:
            pdf_bytes = None
        # 2) MS Word COM (pywin32)
        if pdf_bytes is None:
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(path)
                out_pdf = os.path.splitext(path)[0] + ".pdf"
                doc.SaveAs(out_pdf, FileFormat=17)
                doc.Close()
                word.Quit()
                with open(out_pdf, "rb") as f:
                    pdf_bytes = f.read()
                try:
                    os.remove(out_pdf)
                except Exception:
                    pass
            except Exception:
                pdf_bytes = None
        # 3) LibreOffice (soffice)
        if pdf_bytes is None:
            try:
                soffice = shutil.which("soffice")
                if soffice:
                    out_dir = os.path.dirname(path)
                    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, path], check=True)
                    out_pdf = os.path.splitext(path)[0] + ".pdf"
                    if os.path.exists(out_pdf):
                        with open(out_pdf, "rb") as f:
                            pdf_bytes = f.read()
                        try:
                            os.remove(out_pdf)
                        except Exception:
                            pass
            except Exception:
                pdf_bytes = None
        # 4) Pandoc
        if pdf_bytes is None:
            try:
                pandoc = shutil.which("pandoc")
                if pandoc:
                    out_pdf = os.path.splitext(path)[0] + ".pdf"
                    subprocess.run([pandoc, path, "-o", out_pdf], check=True)
                    with open(out_pdf, "rb") as f:
                        pdf_bytes = f.read()
                    try:
                        os.remove(out_pdf)
                    except Exception:
                        pass
            except Exception:
                pdf_bytes = None
        # 5) python-docx + reportlab (文本版)
        if pdf_bytes is None:
            try:
                import docx
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import mm
                document = docx.Document(path)
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=A4)
                width, height = A4
                y = height - 20 * mm
                for para in document.paragraphs:
                    text = para.text
                    if not text:
                        y -= 8 * mm
                        continue
                    c.drawString(20 * mm, y, text)
                    y -= 10 * mm
                    if y < 20 * mm:
                        c.showPage()
                        y = height - 20 * mm
                c.save()
                pdf_bytes = buffer.getvalue()
            except Exception:
                pdf_bytes = None
        if pdf_bytes is None:
            # 6) Minimal pure-Python fallback: parse DOCX XML to plain text and render a basic PDF
            try:
                def _extract_docx_plain_text(docx_path: str):
                    lines = []
                    with zipfile.ZipFile(docx_path) as z:
                        with z.open("word/document.xml") as f:
                            xml = f.read()
                    root = ET.fromstring(xml)
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for p in root.findall('.//w:p', ns):
                        texts = []
                        for t in p.findall('.//w:t', ns):
                            texts.append(t.text or "")
                        lines.append("".join(texts))
                    return lines
                def _escape_pdf_text(s: str):
                    return s.replace("(", "\\(").replace(")", "\\)").replace("\\", "\\\\")
                def _make_pdf_from_lines(lines):
                    buf = io.BytesIO()
                    parts = []
                    def w(x):
                        parts.append(x.encode('latin-1', errors='replace'))
                    w("%PDF-1.4\n")
                    offsets = []
                    def obj(n, content):
                        offsets.append(sum(len(p) for p in parts))
                        w(f"{n} 0 obj\n")
                        w(content)
                        w("\nendobj\n")
                    # Objects
                    obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
                    obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
                    obj(4, "<< /Type /Font /Subtype /Type1 /Name /F1 /BaseFont /Helvetica >>")
                    # Content stream
                    leading = 16
                    y = 800
                    content = ["BT", "/F1 12 Tf", f"1 0 0 1 50 {y} Tm", f"{leading} TL"]
                    for line in lines[:2000]:
                        txt = _escape_pdf_text(line)
                        content.append(f"({txt}) Tj T*")
                    content.append("ET")
                    stream = "\n".join(content).encode('latin-1', errors='replace')
                    obj(5, f"<< /Length {len(stream)} >>\nstream\n" + stream.decode('latin-1', errors='replace') + "\nendstream")
                    obj(3, "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
                    # xref
                    xref_offset = sum(len(p) for p in parts)
                    w("xref\n")
                    w("0 6\n")
                    w("0000000000 65535 f \n")
                    for off in offsets:
                        w(f"{off:010d} 00000 n \n")
                    w("trailer\n")
                    w("<< /Root 1 0 R /Size 6 >>\n")
                    w("startxref\n")
                    w(f"{xref_offset}\n")
                    w("%%EOF\n")
                    for p in parts:
                        buf.write(p)
                    return buf.getvalue()
                lines = _extract_docx_plain_text(path)
                pdf_bytes = _make_pdf_from_lines(lines)
            except Exception:
                pdf_bytes = None
        if pdf_bytes is None:
            return "Error: Conversion failed (no available converter: docx2pdf, Word COM, LibreOffice, Pandoc, or pure-Python fallback)"
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        aid = _put_artifact(b64)
        out_name = os.path.splitext(filename)[0] + ".pdf"
        return f"Converted successfully. [FILE_ID:{aid}:pdf:{out_name}]"
    except Exception as e:
        return f"Error converting Word to PDF: {str(e)}"
    finally:
        try:
            os.remove(path)
        except Exception:
            pass

@tool
def pdf_to_word_from_upload(filename: str) -> str:
    """Convert an uploaded PDF file to Word (DOCX); returns a downloadable artifact."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        path = os.path.join("uploads", filename)
        if not os.path.exists(path):
            return "Error: File not found in uploads/"
        docx_bytes = None
        # 1) pdf2docx (layout-aware)
        try:
            from pdf2docx import Converter
            out_docx = os.path.splitext(path)[0] + ".docx"
            cv = Converter(path)
            cv.convert(out_docx, start=0, end=None)
            cv.close()
            with open(out_docx, "rb") as f:
                docx_bytes = f.read()
            try:
                os.remove(out_docx)
            except Exception:
                pass
        except Exception:
            docx_bytes = None
        # 2) LibreOffice
        if docx_bytes is None:
            try:
                soffice = shutil.which("soffice")
                if soffice:
                    out_dir = os.path.dirname(path)
                    subprocess.run([soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, path], check=True)
                    out_docx = os.path.splitext(path)[0] + ".docx"
                    if os.path.exists(out_docx):
                        with open(out_docx, "rb") as f:
                            docx_bytes = f.read()
                        try:
                            os.remove(out_docx)
                        except Exception:
                            pass
            except Exception:
                docx_bytes = None
        # 3) pdfminer + python-docx
        if docx_bytes is None:
            try:
                from pdfminer.high_level import extract_text
                import docx
                text = extract_text(path) or ""
                document = docx.Document()
                for line in text.splitlines():
                    document.add_paragraph(line)
                buf = io.BytesIO()
                document.save(buf)
                docx_bytes = buf.getvalue()
            except Exception:
                docx_bytes = None
        # 4) PyPDF2 + python-docx
        if docx_bytes is None:
            try:
                from PyPDF2 import PdfReader
                import docx
                reader = PdfReader(path)
                document = docx.Document()
                for page in reader.pages:
                    content = page.extract_text() or ""
                    for line in (content.splitlines() if content else [""]):
                        document.add_paragraph(line)
                buf = io.BytesIO()
                document.save(buf)
                docx_bytes = buf.getvalue()
            except Exception:
                docx_bytes = None
        if docx_bytes is None:
            return "Error: Conversion failed (pdfminer/PyPDF2 + python-docx not available)"
        b64 = base64.b64encode(docx_bytes).decode("utf-8")
        aid = _put_artifact(b64)
        out_name = os.path.splitext(filename)[0] + ".docx"
        return f"Converted successfully. [FILE_ID:{aid}:docx:{out_name}]"
    except Exception as e:
        return f"Error converting PDF to Word: {str(e)}"
    finally:
        try:
            os.remove(path)
        except Exception:
            pass

@tool
def excel_to_pdf_from_upload(filename: str) -> str:
    """Convert an uploaded Excel file to PDF (tabular rendering); returns a downloadable artifact."""
    try:
        if not _allowed(filename):
            return "Error: File not allowed (not in current session uploads)"
        path = os.path.join("uploads", filename)
        if not os.path.exists(path):
            return "Error: File not found in uploads/"
        pdf_bytes = None
        try:
            df = pd.read_excel(path)
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4)
            data = [list(df.columns)] + df.astype(str).values.tolist()
            tbl = Table(data)
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#999999')),
                ('FONT', (0,0), (-1,-1), 'Helvetica', 9),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ]))
            doc.build([tbl])
            pdf_bytes = buf.getvalue()
        except Exception:
            pdf_bytes = None
        if pdf_bytes is None:
            return "Error: Conversion failed (reportlab not available)"
        b64 = base64.b64encode(pdf_bytes).decode('utf-8')
        aid = _put_artifact(b64)
        out_name = os.path.splitext(filename)[0] + ".pdf"
        return f"Converted successfully. [FILE_ID:{aid}:pdf:{out_name}]"
    except Exception as e:
        return f"Error converting Excel to PDF: {str(e)}"
    finally:
        try:
            os.remove(path)
        except Exception:
            pass
